"""
File exporters for meeting notes: Markdown, Word (.docx) and PDF.

* Markdown / Word handle every language (the viewer supplies the fonts).
* PDF is rendered with the bundled NanumGothic font, which covers Korean +
  Latin well. Scripts the font lacks (e.g. Arabic, full CJK, Cyrillic) may not
  render — recommend Word for those. We never raise on a missing glyph.
"""

from __future__ import annotations

import logging
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos

# fontTools logs every glyph it subsets at INFO; keep it quiet.
logging.getLogger("fontTools").setLevel(logging.WARNING)

_FONT_DIR = Path(__file__).parent.parent / "assets" / "fonts"
_FONT_REGULAR = _FONT_DIR / "NanumGothic.ttf"
_FONT_BOLD = _FONT_DIR / "NanumGothic-Bold.ttf"


def _now() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def build_markdown(title: str, entries: Iterable[dict], summary: str = "") -> bytes:
    lines = [f"# {title}", "", f"_{_now()}_", ""]
    if summary:
        lines += ["## 요약 / Summary", "", summary, ""]
    lines += ["## 전체 기록 / Transcript", ""]
    for e in entries:
        t = e.get("time", "")
        lines.append(f"- **[{t}]** {e.get('source', '')}")
        if e.get("translation"):
            lines.append(f"  - → {e['translation']}")
        if e.get("risk"):
            lines.append(f"  - ⚠️ {e['risk']}")
    return ("\n".join(lines) + "\n").encode("utf-8")


def build_docx(title: str, entries: Iterable[dict], summary: str = "") -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(_now())

    if summary:
        doc.add_heading("요약 / Summary", level=1)
        doc.add_paragraph(summary)

    doc.add_heading("전체 기록 / Transcript", level=1)
    for e in entries:
        p = doc.add_paragraph()
        p.add_run(f"[{e.get('time', '')}] ").bold = True
        p.add_run(e.get("source", ""))
        if e.get("translation"):
            doc.add_paragraph("→ " + e["translation"])
        if e.get("risk"):
            doc.add_paragraph("[!] " + e["risk"])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(title: str, entries: Iterable[dict], summary: str = "") -> bytes:
    pdf = FPDF()
    pdf.add_font("Nanum", "", str(_FONT_REGULAR))
    pdf.add_font("Nanum", "B", str(_FONT_BOLD))
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Always return the cursor to the left margin on the next line, otherwise
    # fpdf2 leaves x at the right edge and the following cell has zero width.
    def cell(text: str, h: float = 6) -> None:
        pdf.multi_cell(0, h, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Nanum", "B", 16)
    cell(title, 9)
    pdf.set_font("Nanum", "", 9)
    pdf.set_text_color(120, 120, 120)
    cell(_now())
    pdf.set_text_color(0, 0, 0)
    pdf.ln(2)

    if summary:
        pdf.set_font("Nanum", "B", 13)
        cell("요약 / Summary", 8)
        pdf.set_font("Nanum", "", 11)
        cell(summary)
        pdf.ln(2)

    pdf.set_font("Nanum", "B", 13)
    cell("전체 기록 / Transcript", 8)
    for e in entries:
        pdf.set_font("Nanum", "B", 10)
        cell(f"[{e.get('time', '')}] {e.get('source', '')}")
        if e.get("translation"):
            pdf.set_font("Nanum", "", 11)
            cell("  → " + e["translation"])
        if e.get("risk"):
            pdf.set_font("Nanum", "", 9)
            pdf.set_text_color(180, 60, 60)
            cell("  [!] " + e["risk"])
            pdf.set_text_color(0, 0, 0)
        pdf.ln(1)

    out = pdf.output()
    return bytes(out)


EXPORTERS = {
    "md": (build_markdown, "text/markdown; charset=utf-8", "md"),
    "docx": (
        build_docx,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ),
    "pdf": (build_pdf, "application/pdf", "pdf"),
}
